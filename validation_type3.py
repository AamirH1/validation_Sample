import sys
import difflib
import re
import json
import argparse
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple, Union
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
import openai  # pip install openai

# For PDF processing - you'll need to install: pip install PyMuPDF
try:
    import fitz  # PyMuPDF for PDF processing
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("[Warning] PyMuPDF not installed. PDF processing will be limited.")


# File processing constants
CHUNK_PAGE_SIZE = 5  # Number of pages per chunk
WORDS_PER_PAGE_ESTIMATE = 250  # Rough estimate for DOCX chunking
PDF_EXTENSIONS = ['.pdf']
DOCX_EXTENSIONS = ['.docx']


def extract_table_content(table: Table) -> str:
    """
    Extract content from a Word table as structured text.
    """
    table_text = []
    try:
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_cells.append(cell_text if cell_text else "")
            if any(cell for cell in row_cells):  # Only add rows with content
                table_text.append(" | ".join(row_cells))
        return "\n".join(table_text) if table_text else ""
    except Exception as e:
        print(f"[Warning] Error extracting table content: {e}")
        return ""


def is_likely_heading(paragraph: Paragraph) -> bool:
    """
    Enhanced heading detection using multiple criteria (used for sub-section context).
    """
    text = paragraph.text.strip()
    if not text:
        return False
    
    if paragraph.style and 'heading' in paragraph.style.name.lower():
        return True
    
    if paragraph.runs:
        first_run = paragraph.runs[0]
        is_bold = bool(first_run.bold)  # Coerce None to False
        font_size = first_run.font.size
        is_larger_font = font_size and font_size.pt > 11 if font_size else False
        is_short = len(text) < 100
        is_formatted_title = text.isupper() or text.istitle()
        heading_score = sum(bool(v) for v in [is_bold and is_short, is_larger_font, is_formatted_title, text.endswith(':') and is_short])
        return heading_score >= 2
    
    return False


def extract_text_and_tables_from_docx(docx_path: str, debug_mode: bool = False) -> List[Dict[str, Any]]:
    """
    Extract text and tables from DOCX file and return as chunks (simulating pages).
    Each chunk represents approximately 5 pages worth of content.
    """
    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"[Error] Could not open DOCX '{docx_path}': {e}")
        return []

    all_paragraphs = []
    all_tables = []

    # Extract all paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            all_paragraphs.append(para.text.strip())

    # Extract all tables
    for table in doc.tables:
        table_content = extract_table_content(table)
        if table_content:
            all_tables.append(table_content)

    # Combine all text
    full_text = '\n\n'.join(all_paragraphs)
    
    if not full_text:
        return []

    # Split into chunks based on word count (simulating pages)
    words = full_text.split()
    words_per_chunk = WORDS_PER_PAGE_ESTIMATE * CHUNK_PAGE_SIZE
    
    chunks = []
    chunk_number = 1
    
    for i in range(0, len(words), words_per_chunk):
        chunk_words = words[i:i + words_per_chunk]
        chunk_text = ' '.join(chunk_words)
        
        # Distribute tables across chunks (simple approach)
        chunk_tables = []
        if all_tables:
            tables_per_chunk = max(1, len(all_tables) // max(1, (len(words) // words_per_chunk)))
            start_table = min((chunk_number - 1) * tables_per_chunk, len(all_tables))
            end_table = min(start_table + tables_per_chunk, len(all_tables))
            chunk_tables = all_tables[start_table:end_table]
        
        chunks.append({
            'chunk_number': chunk_number,
            'text': chunk_text,
            'tables': chunk_tables,
            'word_count': len(chunk_words),
            'page_range': f"{(chunk_number-1)*CHUNK_PAGE_SIZE + 1}-{chunk_number*CHUNK_PAGE_SIZE}"
        })
        
        if debug_mode:
            print(f"[Debug] DOCX Chunk {chunk_number}: {len(chunk_words)} words, {len(chunk_tables)} tables")
        
        chunk_number += 1

    return chunks


def extract_text_and_tables_from_pdf(pdf_path: str, debug_mode: bool = False) -> List[Dict[str, Any]]:
    """
    Extract text and tables from PDF file in 5-page chunks.
    Uses PyMuPDF for extraction.
    """
    if not PDF_AVAILABLE:
        print("[Error] PyMuPDF not available. Cannot process PDF files.")
        return []

    try:
        doc = fitz.open(pdf_path)
    except Exception as e:
        print(f"[Error] Could not open PDF '{pdf_path}': {e}")
        return []

    chunks = []
    total_pages = len(doc)
    
    for start_page in range(0, total_pages, CHUNK_PAGE_SIZE):
        end_page = min(start_page + CHUNK_PAGE_SIZE, total_pages)
        chunk_number = (start_page // CHUNK_PAGE_SIZE) + 1
        
        chunk_text = []
        chunk_tables = []
        
        for page_num in range(start_page, end_page):
            page = doc[page_num]
            text = page.get_text()
            if text.strip():
                chunk_text.append(text.strip())
            
            # Simple table detection (PyMuPDF has basic table support)
            try:
                tables = page.find_tables()
                for table in tables:
                    table_data = table.extract()
                    if table_data:
                        table_text = '\n'.join([' | '.join(row) for row in table_data if any(cell for cell in row)])
                        if table_text:
                            chunk_tables.append(table_text)
            except:
                pass  # Table extraction may not be available in all PyMuPDF versions
        
        combined_text = '\n\n'.join(chunk_text)
        
        if combined_text:  # Only add chunks with content
            chunks.append({
                'chunk_number': chunk_number,
                'text': combined_text,
                'tables': chunk_tables,
                'word_count': len(combined_text.split()),
                'page_range': f"{start_page + 1}-{end_page}"
            })
            
            if debug_mode:
                print(f"[Debug] PDF Chunk {chunk_number} (Pages {start_page + 1}-{end_page}): {len(combined_text.split())} words, {len(chunk_tables)} tables")

    doc.close()
    return chunks


def extract_sections_aggregate_subcontent(
    docx_path: str,
    parent_headings: List[str],
    min_content_length: int,
    debug_mode: bool
) -> List[Dict[str, Any]]:
    """
    Enhanced section extraction with table support and better content aggregation.
    Only splits sections using explicit parent section headings—case insensitive, trimmed.
    All paragraph content (including subheadings, bullets, tables, etc.) is grouped under the last parent heading.
    """
    if not Path(docx_path).exists():
        print(f"[Error] File '{docx_path}' not found!")
        return []
    
    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"[Error] Could not open document '{docx_path}': {e}")
        return []
    
    sections = []
    current_heading = None
    current_content = []
    current_tables = []
    normalized_parents = [h.strip().lower() for h in parent_headings]
    
    doc_elements = []
    para_map = {id(p._element): p for p in doc.paragraphs}
    table_map = {id(t._element): t for t in doc.tables}
    
    for element in doc.element.body:
        if element.tag.endswith('p'):
            para = para_map.get(id(element))
            if para: doc_elements.append(('paragraph', para))
        elif element.tag.endswith('tbl'):
            table = table_map.get(id(element))
            if table: doc_elements.append(('table', table))

    for element_type, element in doc_elements:
        if element_type == 'paragraph':
            text = element.text.strip()
            if not text:
                continue
            
            norm_text = text.strip(" *").replace("**", "").lower()
            
            if norm_text in normalized_parents:
                if current_heading is not None:
                    sections.append({
                        "heading": current_heading,
                        "text": "\n".join(current_content).strip(),
                        "tables": current_tables,
                        "content_length": len("\n".join(current_content).strip()),
                        "table_count": len(current_tables)
                    })
                
                current_heading = text.strip(" *").replace("**", "")
                current_content = []
                current_tables = []
                
                if debug_mode:
                    print(f"[Debug] Found parent section: '{current_heading}'")
            else:
                if current_heading:
                    current_content.append(text)
                    if is_likely_heading(element) and len(text) < 100 and debug_mode:
                        print(f"[Debug] Detected subsection in '{current_heading}': '{text[:50]}...'")
        
        elif element_type == 'table':
            if current_heading:
                table_content = extract_table_content(element)
                if table_content:
                    current_tables.append(table_content)
                    current_content.append(f"\n[TABLE]\n{table_content}\n[/TABLE]\n")
                    if debug_mode:
                        print(f"[Debug] Extracted table in '{current_heading}': {len(table_content)} chars")

    if current_heading:
        sections.append({
            "heading": current_heading,
            "text": "\n".join(current_content).strip(),
            "tables": current_tables,
            "content_length": len("\n".join(current_content).strip()),
            "table_count": len(current_tables)
        })

    valid_sections = [s for s in sections if s["content_length"] >= min_content_length]
    if debug_mode:
        for section in sections:
            if section["content_length"] < min_content_length:
                print(f"[Warning] Section '{section['heading']}' has minimal content ({section['content_length']} chars), skipping.")
    
    return valid_sections


def find_section(sections: List[Dict], desired_heading: str, max_fuzzy_matches: int, debug_mode: bool) -> Optional[Dict]:
    """
    Enhanced section finder with improved fuzzy matching and detailed logging.
    """
    if not sections or not desired_heading:
        return None
    
    for s in sections:
        if s['heading'].strip().lower() == desired_heading.strip().lower():
            if debug_mode:
                print(f"[Debug] Exact match found for '{desired_heading}' -> '{s['heading']}'")
            return s
    
    headings = [s['heading'] for s in sections]
    close_matches = difflib.get_close_matches(
        desired_heading.lower(),
        [h.lower() for h in headings],
        n=max_fuzzy_matches,
        cutoff=0.6
    )
    
    if close_matches:
        best_match = close_matches[0]
        idx = [h.lower() for h in headings].index(best_match)
        match_ratio = difflib.SequenceMatcher(None, desired_heading.lower(), best_match).ratio()
        if debug_mode:
            print(f"[Debug] Fuzzy match for '{desired_heading}' -> '{headings[idx]}' (ratio: {match_ratio:.2f})")
        return sections[idx]
    
    if debug_mode:
        print(f"[Debug] No match found for '{desired_heading}'. Available sections: {[s['heading'] for s in sections]}")
    return None


def analyze_content_quality(text: str, tables: List[str]) -> Dict[str, Any]:
    """
    Analyze content quality metrics for better validation.
    """
    total_text = text + "\n".join(tables)
    word_count = len(total_text.split())
    sentence_count = len([s for s in total_text.split('.') if s.strip()])
    has_lists = bool(re.search(r'^\s*[-•*]\s+', total_text, re.MULTILINE))
    has_numbers = bool(re.search(r'\d+', total_text))
    has_technical_terms = bool(re.search(r'[A-Z]{2,}|API|SDK|URL|HTTP', total_text))
    
    return {
        "word_count": word_count,
        "sentence_count": sentence_count,
        "has_lists": has_lists,
        "has_numbers": has_numbers,
        "has_technical_terms": has_technical_terms,
        "table_count": len(tables),
        "quality_score": min(100, word_count * 0.5 + sentence_count * 2 + len(tables) * 10)
    }


def get_openai_client() -> Optional[Union[openai.AzureOpenAI, openai.OpenAI]]:
    """Initializes and returns an OpenAI client from global configuration."""
    api_key = OPENAI_CONFIG.get("api_key")
    endpoint = OPENAI_CONFIG.get("endpoint")
    api_version = OPENAI_CONFIG.get("api_version")

    if not api_key:
        print("[Error] No OpenAI API key provided in config!")
        return None
    
    try:
        if endpoint:
            return openai.AzureOpenAI(api_key=api_key, azure_endpoint=endpoint, api_version=api_version)
        else:
            return openai.OpenAI(api_key=api_key)
    except Exception as e:
        print(f"[Error] Failed to initialize OpenAI client: {e}")
        return None


def validate_sections_with_llm(
    sections: List[Dict[str, Any]],
    headings: List[str],
    config: Dict[str, Any]
) -> Dict[str, str]:
    """
    Enhanced validation with better error handling and content analysis.
    """
    doc_type = config.get("doc_type", "document")
    debug_mode = config.get("debug_mode", False)
    max_fuzzy_matches = config.get("max_fuzzy_matches", 3)
    
    deployment_name = OPENAI_CONFIG.get("deployment_name")

    try:
        prompt_config = config['prompts']['sections_validation']
        system_prompt = prompt_config.get("system_prompt", "You are a helpful assistant.")
        user_prompt_template = prompt_config.get("user_prompt_template")
    except (KeyError, TypeError):
        print("[Error] 'sections_validation' prompts not found or invalid in config.")
        return {h: f"{h}: ERROR - Prompt configuration missing" for h in headings}

    if not user_prompt_template:
        print("[Error] 'user_prompt_template' not found in 'sections_validation' config.")
        return {h: f"{h}: ERROR - Prompt template missing" for h in headings}

    client = get_openai_client()
    if not client:
        return {h: f"{h}: ERROR - Client initialization failed" for h in headings}
    
    result = {}
    
    for i, heading in enumerate(headings):
        print(f"   [Step] Validating section '{heading}' ({i+1}/{len(headings)}) ...")
        
        sect = find_section(sections, heading, max_fuzzy_matches, debug_mode)
        if not sect:
            result[heading] = f"{heading}: Missing content for {heading}."
            print(f"      [Result] Section not found")
            continue
        
        text = sect['text']
        tables = sect.get('tables', [])
        quality_metrics = analyze_content_quality(text, tables)
        
        if debug_mode:
            print(f"      [Debug] Section: '{sect['heading']}'")
            print(f"      [Debug] Content: {quality_metrics['word_count']} words, {len(tables)} tables")
            print(f"      [Debug] Quality score: {quality_metrics['quality_score']:.1f}/100")
        
        if quality_metrics['word_count'] < 10:
            result[heading] = f"{heading}: Missing content for {heading}."
            print(f"      [Result] Insufficient content ({quality_metrics['word_count']} words)")
            continue
        
        prompt = user_prompt_template.format(
            heading=heading,
            doc_type=doc_type,
            text=text[:2000] + ("..." if len(text) > 2000 else ""),
            tables='\n---\n'.join(tables[:3]) + ("..." if len(tables) > 3 else ""),
            **quality_metrics
        )
        
        try:
            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": prompt}
            ]
            response = client.chat.completions.create(
                messages=messages,
                model=deployment_name,
                max_tokens=150,
                temperature=0.0
            )
            response_text = response.choices[0].message.content.strip()
            result[heading] = response_text
            print(f"      [Step] Validation complete for '{heading}'")
            if debug_mode:
                print(f"      [Result] {response_text}")
                
        except Exception as e:
            error_msg = f"{heading}: ERROR - {str(e)[:100]}"
            result[heading] = error_msg
            print(f"      [Error] Validation failed for '{heading}': {e}")
    
    return result


def validate_raw_text_against_headings(
    raw_text: str, headings: List[str], config: Dict[str, Any]
) -> Dict[str, str]:
    """
    Validates a block of raw text against a list of desired section headings using an LLM.
    """
    doc_type = config.get("doc_type", "document")
    
    deployment_name = OPENAI_CONFIG.get("deployment_name")

    try:
        prompt_config = config['prompts']['raw_text_validation']
        system_prompt = prompt_config.get("system_prompt", "You are a helpful assistant.")
        user_prompt_template = prompt_config.get("user_prompt_template")
    except (KeyError, TypeError):
        print("[Error] 'raw_text_validation' prompts not found or invalid in config.")
        return {h: f"{h}: ERROR - Prompt configuration missing" for h in headings}

    if not user_prompt_template:
        print("[Error] 'user_prompt_template' not found in 'raw_text_validation' config.")
        return {h: f"{h}: ERROR - Prompt template missing" for h in headings}

    client = get_openai_client()
    if not client:
        return {h: f"{h}: ERROR - Client initialization failed" for h in headings}
    
    result = {}
    truncated_text = raw_text[:10000] + ("..." if len(raw_text) > 10000 else "")

    for i, heading in enumerate(headings):
        print(f"   [Step] Analyzing raw text for section '{heading}' ({i+1}/{len(headings)}) ...")
        prompt = user_prompt_template.format(heading=heading, doc_type=doc_type, text=truncated_text)
        try:
            messages = [{"role": "system", "content": system_prompt}, {"role": "user", "content": prompt}]
            response = client.chat.completions.create(messages=messages, model=deployment_name, max_tokens=150, temperature=0.0)
            result[heading] = response.choices[0].message.content.strip()
            print(f"      [Step] Analysis complete for '{heading}'")
        except Exception as e:
            result[heading] = f"{heading}: ERROR - {str(e)[:100]}"
            print(f"      [Error] Analysis failed for '{heading}': {e}")
    return result


def validate_file_content_against_headings(
    chunks: List[Dict[str, Any]], 
    headings: List[str], 
    config: Dict[str, Any]
) -> Dict[str, str]:
    """
    NEW: Validates extracted file chunks against desired headings.
    Processes chunks sequentially and aggregates content for each section.
    """
    doc_type = config.get("doc_type", "document")
    debug_mode = config.get("debug_mode", False)
    deployment_name = OPENAI_CONFIG.get("deployment_name")

    try:
        prompt_config = config['prompts']['file_validation']
        system_prompt = prompt_config.get("system_prompt", "You are a helpful assistant.")
        user_prompt_template = prompt_config.get("user_prompt_template")
    except (KeyError, TypeError):
        print("[Error] 'file_validation' prompts not found or invalid in config.")
        return {h: f"{h}: ERROR - Prompt configuration missing" for h in headings}

    if not user_prompt_template:
        print("[Error] 'user_prompt_template' not found in 'file_validation' config.")
        return {h: f"{h}: ERROR - Prompt template missing" for h in headings}

    client = get_openai_client()
    if not client:
        return {h: f"{h}: ERROR - Client initialization failed" for h in headings}

    # Aggregate content across all chunks
    section_content = {heading: {'text_parts': [], 'tables': [], 'total_words': 0} for heading in headings}
    
    print(f"   [Info] Processing {len(chunks)} chunks to extract section content...")
    
    for chunk in chunks:
        chunk_text = chunk['text']
        chunk_tables = chunk.get('tables', [])
        
        if debug_mode:
            print(f"   [Debug] Processing Chunk {chunk['chunk_number']} (Pages {chunk['page_range']}): {chunk['word_count']} words")
        
        # For each heading, extract relevant content from this chunk
        for heading in headings:
            try:
                # Use LLM to extract relevant content for this heading from this chunk
                extraction_prompt = f"""Extract any content from the following text that would be relevant for creating a '{heading}' section in a {doc_type}. 

Text:
{chunk_text[:3000]}

If there is relevant content, return it. If there is no relevant content, return "NO_RELEVANT_CONTENT".
Focus on substantial, detailed information that could be used to populate the '{heading}' section."""

                messages = [
                    {"role": "system", "content": "You are an expert content extractor. Extract only relevant, substantial content for the specified section."},
                    {"role": "user", "content": extraction_prompt}
                ]
                
                response = client.chat.completions.create(
                    messages=messages,
                    model=deployment_name,
                    max_tokens=800,
                    temperature=0.0
                )
                
                extracted_content = response.choices[0].message.content.strip()
                
                if extracted_content != "NO_RELEVANT_CONTENT" and len(extracted_content.split()) > 5:
                    section_content[heading]['text_parts'].append(extracted_content)
                    section_content[heading]['total_words'] += len(extracted_content.split())
                    
                    if debug_mode:
                        print(f"      [Debug] Found content for '{heading}' in chunk {chunk['chunk_number']}: {len(extracted_content.split())} words")
                
                # Add relevant tables
                if chunk_tables:
                    section_content[heading]['tables'].extend(chunk_tables)
                    
            except Exception as e:
                if debug_mode:
                    print(f"      [Warning] Content extraction failed for '{heading}' in chunk {chunk['chunk_number']}: {e}")
                continue
    
    # Now validate each section based on aggregated content
    result = {}
    
    for heading in headings:
        print(f"   [Step] Validating aggregated content for '{heading}'...")
        
        aggregated_text = '\n\n'.join(section_content[heading]['text_parts'])
        section_tables = section_content[heading]['tables']
        total_words = section_content[heading]['total_words']
        
        if debug_mode:
            print(f"      [Debug] '{heading}' - Aggregated: {total_words} words, {len(section_tables)} tables")
        
        if total_words < 20:  # Insufficient content
            result[heading] = f"{heading}: Insufficient content found across all chunks to populate this section."
            continue
        
        # Quality analysis
        quality_metrics = analyze_content_quality(aggregated_text, section_tables)
        
        # Final validation prompt
        final_prompt = user_prompt_template.format(
            heading=heading,
            doc_type=doc_type,
            text=aggregated_text[:2000] + ("..." if len(aggregated_text) > 2000 else ""),
            tables='\n---\n'.join(section_tables[:2]) + ("..." if len(section_tables) > 2 else ""),
            **quality_metrics
        )
        
        try:
            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": final_prompt}
            ]
            
            response = client.chat.completions.create(
                messages=messages,
                model=deployment_name,
                max_tokens=150,
                temperature=0.0
            )
            
            response_text = response.choices[0].message.content.strip()
            result[heading] = response_text
            
            if debug_mode:
                print(f"      [Result] {response_text}")
                
        except Exception as e:
            error_msg = f"{heading}: ERROR - {str(e)[:100]}"
            result[heading] = error_msg
            print(f"      [Error] Final validation failed for '{heading}': {e}")
    
    return result


def load_config(config_path: str) -> Optional[Dict]:
    """Loads the configuration from a JSON file."""
    path = Path(config_path)
    if not path.is_file():
        print(f"[Error] Configuration file not found at '{config_path}'")
        return None
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except json.JSONDecodeError as e:
        print(f"[Error] Invalid JSON in configuration file '{config_path}': {e}")
        return None
    except Exception as e:
        print(f"[Error] Could not read configuration file '{config_path}': {e}")
        return None


def run_section_validation(config: Dict[str, Any]):
    """Runs the main logic for validating a structured DOCX file."""
    print("=" * 80)
    print("DOCX SECTION VALIDATOR (SECTIONS MODE)")
    print("=" * 80)

    docx_path = config.get("input_docx_filename")
    if not docx_path:
        print("[Error] 'input_docx_filename' not specified in the profile for SECTIONS mode. Exiting.")
        return

    print(f"[Start] Using DOCX file: {docx_path}")
    if not Path(docx_path).is_file():
        print(f"[Error] File '{docx_path}' not found! Exiting.")
        return

    print("\n[Step 1] Extracting parent sections with enhanced content aggregation...")
    sections = extract_sections_aggregate_subcontent(
        docx_path, config["parent_section_headings"], config["min_content_length"], config["debug_mode"])
    
    if not sections:
        print("[Error] No valid sections found. Check your document structure and heading definitions.")
        return

    print(f"\n[Step 2] Successfully extracted {len(sections)} sections:")
    for s in sections:
        quality = analyze_content_quality(s.get('text', ''), s.get('tables', []))
        print(f"   [Section] '{s['heading']}'")
        print(f"     - Content: {s.get('content_length', 0):,} chars ({quality['word_count']} words)")
        print(f"     - Tables: {s.get('table_count', 0)}")
        print(f"     - Quality: {quality['quality_score']:.1f}/100")

    mandatory_headings = config["mandatory_section_headings"]
    print("\n[Step 3] Validating ONLY the following mandatory sections:")
    for h in mandatory_headings:
        print(f"   [Validate] {h}")

    print("\n[Step 4] Validating content relevance with OpenAI/Azure ...")
    validation_results = validate_sections_with_llm(sections, mandatory_headings, config)

    print("\n" + "=" * 60)
    print("VALIDATION REPORT")
    print("=" * 60)
    for heading, verdict in validation_results.items():
        status_final = "Pass:" if any(k in verdict.lower() for k in ["comprehensive and relevant","sufficient", "incomplete", "too generic"]) else "Fail:"
        print(f"{status_final} {verdict}")
    print("=" * 60)


def run_raw_text_validation(config: Dict[str, Any]):
    """Runs the new logic for validating a raw text file."""
    print("=" * 80)
    print("RAW TEXT SUITABILITY VALIDATOR (RAW_TEXT MODE)")
    print("=" * 80)
    
    raw_text_path_str = config.get("raw_text_input_filename")
    if not raw_text_path_str:
        print("[Error] 'raw_text_input_filename' not specified in the profile for RAW_TEXT mode. Exiting.")
        return

    raw_text_path = Path(raw_text_path_str)
    if not raw_text_path.is_file():
        print(f"[Error] Raw input file '{raw_text_path}' not found! Exiting.")
        return

    print(f"[Start] Processing raw input file: {raw_text_path}")
    raw_text = ""
    try:
        if raw_text_path.suffix.lower() == '.docx':
            doc = Document(raw_text_path)
            raw_text = "\n".join([p.text for p in doc.paragraphs])
            print(f"  (File is a .docx, extracted text content)")
        else:
            raw_text = raw_text_path.read_text(encoding='utf-8')
            print(f"  (File is a plain text file, read content)")
    except Exception as e:
        print(f"[Error] Could not read content from '{raw_text_path}': {e}")
        return

    if len(raw_text.strip()) < config.get("min_content_length", 50):
        print(f"[Warning] Raw input file has very little content ({len(raw_text.strip())} chars).")

    mandatory_headings = config["mandatory_section_headings"]
    print("\n[Step 1] Checking raw text suitability for mandatory sections:")
    for h in mandatory_headings:
        print(f"   [Target] {h}")

    print("\n[Step 2] Running AI-powered content analysis...")
    validation_results = validate_raw_text_against_headings(raw_text, mandatory_headings, config)

    print("\n" + "=" * 60)
    print("RAW TEXT SUITABILITY REPORT")
    print("=" * 60)
    for heading, verdict in validation_results.items():
        status_final = "Pass:" if "sufficient" in verdict.lower() else "Pass:" if "incomplete" in verdict.lower() else "Fail:"
        print(f"{status_final} {verdict}")
    print("=" * 60)


def run_file_validation(config: Dict[str, Any]):
    """NEW: Runs file validation for both PDF and DOCX files with chunking."""
    print("=" * 80)
    print("FILE CONTENT VALIDATOR (FILE MODE)")
    print("=" * 80)
    
    file_path_str = config.get("input_file_path")
    if not file_path_str:
        print("[Error] 'input_file_path' not specified in the profile for FILE mode. Exiting.")
        return

    file_path = Path(file_path_str)
    if not file_path.is_file():
        print(f"[Error] Input file '{file_path}' not found! Exiting.")
        return

    file_extension = file_path.suffix.lower()
    supported_extensions = config.get("supported_file_types", [".docx", ".pdf"])
    
    if file_extension not in supported_extensions:
        print(f"[Error] File type '{file_extension}' not supported. Supported types: {supported_extensions}")
        return

    print(f"[Start] Processing {file_extension.upper()} file: {file_path}")
    print(f"[Config] Chunk size: {CHUNK_PAGE_SIZE} pages per chunk")

    # Extract content in chunks
    print("\n[Step 1] Extracting content in chunks...")
    if file_extension in DOCX_EXTENSIONS:
        chunks = extract_text_and_tables_from_docx(str(file_path), config.get("debug_mode", False))
    elif file_extension in PDF_EXTENSIONS:
        chunks = extract_text_and_tables_from_pdf(str(file_path), config.get("debug_mode", False))
    else:
        print(f"[Error] Unsupported file extension: {file_extension}")
        return

    if not chunks:
        print("[Error] No content could be extracted from the file.")
        return

    print(f"[Step 2] Successfully extracted {len(chunks)} chunks:")
    total_words = sum(chunk['word_count'] for chunk in chunks)
    total_tables = sum(len(chunk['tables']) for chunk in chunks)
    
    for chunk in chunks:
        print(f"   [Chunk {chunk['chunk_number']}] Pages {chunk['page_range']}: {chunk['word_count']} words, {len(chunk['tables'])} tables")
    
    print(f"   [Total] {total_words} words, {total_tables} tables across all chunks")

    mandatory_headings = config["mandatory_section_headings"]
    print(f"\n[Step 3] Analyzing content for {len(mandatory_headings)} mandatory sections:")
    for h in mandatory_headings:
        print(f"   [Target] {h}")

    print("\n[Step 4] Processing chunks and validating section content...")
    validation_results = validate_file_content_against_headings(chunks, mandatory_headings, config)

    print("\n" + "=" * 60)
    print("FILE CONTENT VALIDATION REPORT")
    print("=" * 60)
    for heading, verdict in validation_results.items():
        status_final = "Pass:" if any(k in verdict.lower() for k in ["comprehensive", "sufficient", "detailed enough"]) else "Fail:"
        print(f"{status_final} {verdict}")
    print("=" * 60)


def main():
    parser = argparse.ArgumentParser(
        description="Document Validator using LLMs with support for SECTIONS, RAW_TEXT, and FILE validation modes.",
        formatter_class=argparse.RawTextHelpFormatter
    )
    parser.add_argument(
        "--config_file",
        default="configs_UG3.json",
        help="Path to the configuration JSON file."
    )
    parser.add_argument(
        "--debug",
        action="store_true",
        help="Enable debug mode for verbose logging."
    )
    args = parser.parse_args()

    config = load_config(args.config_file)
    if not config:
        return

    active_profile_name = config.get("active_profile")
    available_profiles = list(config.get("validation_profiles", {}).keys())

    if not active_profile_name:
        print("Error: 'active_profile' key not found or is empty in the configuration file.")
        if available_profiles:
            print("\nAvailable profiles are:")
            for profile in available_profiles:
                print(f"  - {profile}")
            print(f"\nPlease add '\"active_profile\": \"<profile_name>\"' to '{args.config_file}'.")
        else:
            print(f"\nNo validation profiles found in '{args.config_file}'.")
        return

    profile_config = config.get("validation_profiles", {}).get(active_profile_name)
    if not profile_config:
        print(f"[Error] Profile '{active_profile_name}' (from 'active_profile') not found in '{args.config_file}'.")
        if available_profiles:
            print(f"Available profiles: {available_profiles}")
        return
    
    # Command-line --debug flag overrides the config setting
    profile_config['debug_mode'] = args.debug or profile_config.get('debug_mode', False)

    validation_mode = profile_config.get("validation_mode")
    if validation_mode == "SECTIONS":
        run_section_validation(profile_config)
    elif validation_mode == "RAW_TEXT":
        run_raw_text_validation(profile_config)
    elif validation_mode == "FILE":  # NEW MODE
        run_file_validation(profile_config)
    else:
        print(f"[Error] Invalid 'validation_mode': '{validation_mode}' in profile '{active_profile_name}'. Must be 'SECTIONS', 'RAW_TEXT', or 'FILE'.")


if __name__ == "__main__":
    main()
