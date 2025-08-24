import sys
import difflib
import json
import argparse
from pathlib import Path
from typing import List, Dict, Any, Optional
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
import openai  # pip install openai
 

# Global debug mode, will be set from config
DEBUG_MODE = False

def load_config(config_path: str) -> Optional[Dict[str, Any]]:
    """Loads the JSON configuration file."""
    path = Path(config_path)
    if not path.is_file():
        print(f"[Error] Configuration file not found at: {path}")
        return None
    try:
        with open(path, 'r', encoding='utf-8') as f:
            config = json.load(f)
        print(f"[Info] Configuration loaded successfully from '{config_path}'")
        return config
    except json.JSONDecodeError as e:
        print(f"[Error] Invalid JSON in configuration file '{path}': {e}")
        return None
    except Exception as e:
        print(f"[Error] Could not read configuration file '{path}': {e}")
        return None


def extract_table_content(table: Table) -> str:
    """
    Extract content from a Word table as structured text.
    """
    table_text = []
    try:
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_cells.append(cell_text if cell_text else "")
            if any(cell for cell in row_cells):  # Only add rows with content
                table_text.append(" | ".join(row_cells))
        return "\n".join(table_text) if table_text else ""
    except Exception as e:
        if DEBUG_MODE:
            print(f"[Warning] Error extracting table content: {e}")
        return ""


def is_likely_heading(paragraph: Paragraph) -> float:
    """
    Determine if a paragraph is likely a heading and return confidence score (0-1).
    Uses multiple criteria for better accuracy.
    """
    text = paragraph.text.strip()
    if not text or len(text) > 200:  # Very long text unlikely to be heading
        return 0.0
    
    confidence = 0.0
    
    # Check style-based headings (highest confidence)
    if paragraph.style and 'heading' in paragraph.style.name.lower():
        confidence += 0.8
    
    # Check formatting-based criteria
    if paragraph.runs:
        first_run = paragraph.runs[0]
        
        # Bold text
        if first_run.bold:
            confidence += 0.3
        
        # Font size (larger than normal)
        font_size = first_run.font.size
        if font_size and hasattr(font_size, 'pt') and font_size.pt > 12:
            confidence += 0.2
        
        # All caps or title case
        if text.isupper() or text.istitle():
            confidence += 0.2
    
    # Text characteristics
    if len(text) < 100:  # Short text more likely to be heading
        confidence += 0.1
    
    if text.endswith(':'):  # Ends with colon
        confidence += 0.1
    
    # Structural indicators
    if any(keyword in text.lower() for keyword in ['features', 'overview', 'introduction', 'getting started', 'setup', 'configuration']):
        confidence += 0.2
    
    return min(1.0, confidence)


def auto_detect_headings(docx_path: str, min_heading_confidence: float) -> List[str]:
    """
    Automatically detect headings in the document using multiple criteria.
    Returns a list of detected heading texts.
    """
    if not Path(docx_path).exists():
        print(f"[Error] File '{docx_path}' not found!")
        return []
    
    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"[Error] Could not open document '{docx_path}': {e}")
        return []
    
    detected_headings = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        confidence = is_likely_heading(para)
        
        if confidence >= min_heading_confidence:
            detected_headings.append(text)
            if DEBUG_MODE:
                print(f"[Debug] Detected heading: '{text}' (confidence: {confidence:.2f})")
    
    if DEBUG_MODE:
        print(f"[Debug] Total headings detected: {len(detected_headings)}")
    
    return detected_headings


def extract_sections_dynamic(docx_path: str, min_heading_confidence: float, min_content_length: int) -> List[Dict[str, Any]]:
    """
    Dynamically extract sections based on detected headings.
    Includes proper table extraction and content aggregation.
    """
    if not Path(docx_path).exists():
        print(f"[Error] File '{docx_path}' not found!")
        return []
    
    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"[Error] Could not open document '{docx_path}': {e}")
        return []
    
    # First, detect all potential headings
    detected_headings = auto_detect_headings(docx_path, min_heading_confidence)
    
    if not detected_headings:
        print("[Warning] No headings detected. Document may not have proper heading structure.")
        return []
    
    sections = []
    current_heading = None
    current_content = []
    current_tables = []
    
    # Create a mapping of document elements (paragraphs and tables) in order
    doc_elements = []
    
    # Get all paragraphs first
    paragraph_map = {id(p._element): p for p in doc.paragraphs}
    table_map = {id(t._element): t for t in doc.tables}
    
    # Process document elements in order
    for element in doc.element.body:
        if element.tag.endswith('p'):  # Paragraph
            para = paragraph_map.get(id(element))
            if para:
                doc_elements.append(('paragraph', para))
        elif element.tag.endswith('tbl'):  # Table
            table = table_map.get(id(element))
            if table:
                doc_elements.append(('table', table))
    
    # Process elements in document order
    for element_type, element in doc_elements:
        if element_type == 'paragraph':
            text = element.text.strip()
            if not text:
                continue
            
            # Check if this paragraph is a detected heading
            if text in detected_headings:
                # Save previous section if exists
                if current_heading is not None:
                    section_text = "\n".join(current_content).strip()
                    if len(section_text) >= min_content_length or current_tables:
                        sections.append({
                            "heading": current_heading,
                            "text": section_text,
                            "tables": current_tables,
                            "content_length": len(section_text),
                            "table_count": len(current_tables)
                        })
                    elif DEBUG_MODE:
                        print(f"[Debug] Skipping section '{current_heading}' - insufficient content ({len(section_text)} chars, min: {min_content_length})")
                
                # Start new section
                current_heading = text
                current_content = []
                current_tables = []
                
                if DEBUG_MODE:
                    print(f"[Debug] Starting new section: '{current_heading}'")
            else:
                # Add content to current section
                if current_heading is not None:
                    current_content.append(text)
        
        elif element_type == 'table':
            if current_heading is not None:
                table_content = extract_table_content(element)
                if table_content:
                    current_tables.append(table_content)
                    # Also add table content to text for comprehensive analysis
                    current_content.append(f"[TABLE]\n{table_content}\n[/TABLE]")
                    
                    if DEBUG_MODE:
                        print(f"[Debug] Added table to section '{current_heading}': {len(table_content)} chars")
    
    # Don't forget the last section
    if current_heading is not None:
        section_text = "\n".join(current_content).strip()
        if len(section_text) >= min_content_length or current_tables:
            sections.append({
                "heading": current_heading,
                "text": section_text,
                "tables": current_tables,
                "content_length": len(section_text),
                "table_count": len(current_tables)
            })
    
    return sections


def find_section(sections: List[Dict], desired_heading: str) -> Optional[Dict]:
    """
    Enhanced section finder with improved fuzzy matching and detailed logging.
    """
    if not sections or not desired_heading:
        return None
    
    # Try exact case-insensitive match first
    for s in sections:
        if s['heading'].strip().lower() == desired_heading.strip().lower():
            if DEBUG_MODE:
                print(f"[Debug] Exact match found for '{desired_heading}' -> '{s['heading']}'")
            return s
    
    # Try fuzzy matching
    headings = [s['heading'] for s in sections]
    close_matches = difflib.get_close_matches(
        desired_heading.lower(),
        [h.lower() for h in headings],
        n=3,
        cutoff=0.6
    )
    
    if close_matches:
        best_match = close_matches[0]
        idx = [h.lower() for h in headings].index(best_match)
        match_ratio = difflib.SequenceMatcher(None, desired_heading.lower(), best_match).ratio()
        
        if DEBUG_MODE:
            print(f"[Debug] Fuzzy match for '{desired_heading}' -> '{headings[idx]}' (ratio: {match_ratio:.2f})")
        
        return sections[idx]
    
    if DEBUG_MODE:
        print(f"[Debug] No match found for '{desired_heading}'. Available sections: {[s['heading'] for s in sections]}")
    
    return None


def validate_sections_with_llm(
    sections: List[Dict[str, Any]],
    headings: List[str],
    config: Dict[str, Any]
) -> Dict[str, str]:
    """
    Enhanced validation with better error handling and content analysis.
    """
    llm_api_key = OPENAI_API_KEY
    deployment_name = DEPLOYMENT_NAME
    endpoint = ENDPOINT
    api_version = API_VERSION
    doc_type = config.get("doc_type", "document")

    try:
        prompt_config = config['prompts']['sections_validation']
        prompt_template = prompt_config['user_prompt_template']
        system_prompt = prompt_config['system_prompt']
    except KeyError as e:
        print(f"[Error] Missing prompt configuration for 'sections_validation': {e}")
        return {h: f"{h}: ERROR - Invalid prompt configuration" for h in headings}

    if not llm_api_key:
        print("[Error] No OpenAI API key provided!")
        return {h: f"{h}: ERROR - No API key" for h in headings}
    
    try:
        if endpoint:
            client = openai.AzureOpenAI(api_key=llm_api_key, azure_endpoint=endpoint, api_version=api_version)
        else:
            client = openai.OpenAI(api_key=llm_api_key)
    except Exception as e:
        print(f"[Error] Failed to initialize OpenAI client: {e}")
        return {h: f"{h}: ERROR - Client initialization failed" for h in headings}
    
    result = {}
    
    for i, heading in enumerate(headings):
        print(f"   [Step] Validating section '{heading}' ({i+1}/{len(headings)}) ...")
        
        sect = find_section(sections, heading)
        if not sect:
            result[heading] = f"{heading}: Missing content for {heading}."
            print(f"      [Result] Section not found")
            continue
        
        text = sect['text']
        tables = sect.get('tables', [])
        content_length = sect.get('content_length', len(text))
        table_count = sect.get('table_count', len(tables))
        
        if DEBUG_MODE:
            print(f"      [Debug] Found section: '{sect['heading']}'")
            print(f"      [Debug] Content: {content_length} chars, {table_count} tables")
        
        # Skip validation for obviously empty content
        if content_length < 10:
            result[heading] = f"{heading}: Missing content for {heading}."
            print(f"      [Result] Insufficient content ({content_length} chars)")
            continue
        
        # Prepare prompt
        prompt = prompt_template.format(
            heading=heading,
            doc_type=doc_type,
            text=text[:2000] + ("..." if len(text) > 2000 else ""),  # Truncate very long text
            tables='\n---\n'.join(tables[:2]) + ("..." if len(tables) > 2 else ""),  # Limit tables
            content_length=content_length,
            table_count=table_count
        )
        
        try:
            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": prompt}
            ]
            
            response = client.chat.completions.create(
                messages=messages,
                model=deployment_name,
                max_tokens=150,
                temperature=0.0
            )
            
            response_text = response.choices[0].message.content.strip()
            result[heading] = response_text
            
            print(f"      [Step] Validation complete for '{heading}'")
            
        except Exception as e:
            error_msg = f"{heading}: ERROR - {str(e)[:100]}"
            result[heading] = error_msg
            print(f"      [Error] Validation failed for '{heading}': {e}")
    
    return result


def validate_raw_text_against_headings(
    raw_text: str,
    headings: List[str],
    config: Dict[str, Any]
) -> Dict[str, str]:
    """
    Validates a block of raw text against a list of desired section headings using an LLM.
    """
    llm_api_key = OPENAI_API_KEY
    deployment_name = DEPLOYMENT_NAME
    endpoint = ENDPOINT
    api_version = API_VERSION
    doc_type = config.get("doc_type", "document")

    try:
        prompt_config = config['prompts']['raw_text_validation']
        prompt_template = prompt_config['user_prompt_template']
        system_prompt = prompt_config['system_prompt']
    except KeyError as e:
        print(f"[Error] Missing prompt configuration for 'raw_text_validation': {e}")
        return {h: f"{h}: ERROR - Invalid prompt configuration" for h in headings}

    if not llm_api_key:
        print("[Error] No OpenAI API key provided!")
        return {h: f"{h}: ERROR - No API key" for h in headings}
    
    try:
        if endpoint:
            client = openai.AzureOpenAI(api_key=llm_api_key, azure_endpoint=endpoint, api_version=api_version)
        else:
            client = openai.OpenAI(api_key=llm_api_key)
    except Exception as e:
        print(f"[Error] Failed to initialize OpenAI client: {e}")
        return {h: f"{h}: ERROR - Client initialization failed" for h in headings}
    
    result = {}
    truncated_text = raw_text[:8000] + ("..." if len(raw_text) > 8000 else "")

    for i, heading in enumerate(headings):
        print(f"   [Step] Analyzing raw text for section '{heading}' ({i+1}/{len(headings)}) ...")
        prompt = prompt_template.format(heading=heading, doc_type=doc_type, text=truncated_text)
        try:
            messages = [{"role": "system", "content": system_prompt}, {"role": "user", "content": prompt}]
            response = client.chat.completions.create(messages=messages, model=deployment_name, max_tokens=150, temperature=0.0)
            result[heading] = response.choices[0].message.content.strip()
            print(f"      [Step] Analysis complete for '{heading}'")
        except Exception as e:
            result[heading] = f"{heading}: ERROR - {str(e)[:100]}"
            print(f"      [Error] Analysis failed for '{heading}': {e}")
    return result


def run_section_validation(config: Dict[str, Any]):
    """The original functionality of validating sections in a DOCX."""
    print(f"\n[Mode] Running in SECTIONS validation mode (dynamic heading detection).")
    
    docx_path = config.get("input_docx_filename")
    if not docx_path:
        print("[Error] 'input_docx_filename' not specified in config for SECTIONS mode.")
        return

    print(f"\n[Start] Processing document: {docx_path}")
    
    if not Path(docx_path).is_file():
        print(f"[Error] File '{docx_path}' not found! Exiting.")
        return
    
    min_heading_confidence = config.get("min_heading_confidence", 0.6)
    min_content_length = config.get("min_content_length", 20)

    print(f"[Config] Debug mode: {'ON' if DEBUG_MODE else 'OFF'}")
    print(f"[Config] Minimum heading confidence: {min_heading_confidence}")
    print(f"[Config] Minimum content length: {min_content_length} characters")

    print("\n[Step 1] Dynamically detecting document structure...")
    sections = extract_sections_dynamic(docx_path, min_heading_confidence, min_content_length)
    
    if not sections:
        print("[Error] No valid sections found. Check your document structure.")
        return

    print(f"\n[Step 2] Successfully extracted {len(sections)} sections:")
    for s in sections:
        print(f"   [Section] '{s['heading']}'")
        print(f"     - Content: {s.get('content_length', 0):,} chars")
        print(f"     - Tables: {s.get('table_count', 0)}")

    if DEBUG_MODE and sections:
        print(f"\n[DEBUG] Sample content from sections:")
        for i, s in enumerate(sections[:3]):
            sample_text = s.get('text', '')
            print(f"   Section {i+1} '{s['heading']}': {sample_text[:100]}{'...' if len(sample_text) > 100 else ''}")

    detected_headings = [s['heading'] for s in sections]
    
    print(f"\n[Step 3] Validating all {len(detected_headings)} detected sections:")
    for h in detected_headings:
        print(f"   [Target] {h}")

    print("\n[Step 4] Running AI-powered content validation...")
    validation_results = validate_sections_with_llm(
        sections=sections, headings=detected_headings, config=config
    )

    print("\n" + "=" * 80)
    print("SECTION VALIDATION RESULTS")
    print("=" * 80)
    
    for heading, verdict in validation_results.items():
        #status_final = "Pass:" if "comprehensive and relevant" in verdict.lower() else "Pass:" if "available" in verdict.lower() else "Fail:"
        status_final = "Pass:" if any(k in verdict.lower() for k in ["sufficient", "incomplete", "too generic"]) else "Fail:"

        print(f"{status_final} {verdict}")
    
    total_sections = len(validation_results)
    comprehensive = sum(1 for v in validation_results.values() if "comprehensive and relevant" in v.lower())
    partial = sum(1 for v in validation_results.values() if "available but" in v.lower())
    missing = sum(1 for v in validation_results.values() if "missing" in v.lower() or "error" in v.lower())
    
    print(f"\n" + "=" * 40)
    print("SUMMARY STATISTICS")
    print("=" * 40)
    print(f"Total sections: {total_sections}")
    print(f"Comprehensive: {comprehensive} ({comprehensive/max(1,total_sections)*100:.1f}%)")
    print(f"Partial/Issues: {partial} ({partial/max(1,total_sections)*100:.1f}%)")
    print(f"Missing/Errors: {missing} ({missing/max(1,total_sections)*100:.1f}%)")
    
    print("\n" + "=" * 80)


def run_raw_text_validation(config: Dict[str, Any]):
    """New functionality to validate raw text against mandatory headings."""
    print(f"\n[Mode] Running in RAW_TEXT validation mode.")
    
    raw_text_filename = config.get("raw_text_input_filename")
    if not raw_text_filename:
        print("[Error] 'raw_text_input_filename' not specified in config for RAW_TEXT mode.")
        return

    raw_text_path = Path(raw_text_filename)
    if not raw_text_path.is_file():
        print(f"[Error] Raw input file '{raw_text_path}' not found! Exiting.")
        return

    print(f"[Start] Processing raw input file: {raw_text_path}")
    raw_text = ""
    try:
        if raw_text_path.suffix.lower() == '.docx':
            doc = Document(raw_text_path)
            raw_text = "\n".join([p.text for p in doc.paragraphs])
            print(f"  (File is a .docx, extracted text content)")
        else: # Assume .txt or other plain text
            raw_text = raw_text_path.read_text(encoding='utf-8')
            print(f"  (File is a plain text file, read content)")
    except Exception as e:
        print(f"[Error] Could not read content from '{raw_text_path}': {e}")
        return

    min_content_length = config.get("min_content_length", 20)
    if len(raw_text.strip()) < min_content_length:
        print(f"[Warning] Raw input file has very little content ({len(raw_text.strip())} chars).")

    mandatory_headings = config.get("mandatory_section_headings", [])
    if not mandatory_headings:
        print("[Warning] 'mandatory_section_headings' is empty in config for RAW_TEXT mode.")
        return

    print("\n[Step 1] Validating raw text against mandatory sections:")
    for h in mandatory_headings:
        print(f"   [Target] {h}")

    print("\n[Step 2] Running AI-powered content analysis...")
    validation_results = validate_raw_text_against_headings(
        raw_text=raw_text, headings=mandatory_headings, config=config
    )

    print("\n" + "=" * 80)
    print("RAW TEXT SUITABILITY REPORT")
    print("=" * 80)
    
    for heading, verdict in validation_results.items():
        status_final = "Pass:" if any(k in verdict.lower() for k in ["sufficient", "incomplete", "too generic"]) else "Fail:"
        print(f"{status_final} {verdict}")
    
    print("\n" + "=" * 80)


def main():
    parser = argparse.ArgumentParser(description="Enhanced DOCX Section Validator.")
    parser.add_argument(
        "config_file",
        type=str,
        help="Path to the JSON configuration file.",
        nargs='?',
        default="configs_UG1.json"
    )
    args = parser.parse_args()

    print("=" * 80)
    print("ENHANCED DOCX SECTION VALIDATOR")
    print("=" * 80)
    
    config = load_config(args.config_file)
    if not config:
        sys.exit(1)

    active_profile_name = config.get("active_profile")
    if not active_profile_name:
        print("[Error] 'active_profile' key not found in the configuration file.")
        sys.exit(1)

    profile_config = config.get("validation_profiles", {}).get(active_profile_name)
    if not profile_config:
        print(f"[Error] Profile '{active_profile_name}' not found in the configuration file.")
        available_profiles = list(config.get("validation_profiles", {}).keys())
        if available_profiles:
            print(f"Available profiles: {available_profiles}")
        sys.exit(1)

    # Set global debug mode from the active profile
    global DEBUG_MODE
    DEBUG_MODE = profile_config.get("debug_mode", False)

    print(f"\n[Info] Using active profile: '{active_profile_name}'")
    print(f"[Info] Document type: '{profile_config.get('doc_type', 'N/A')}'")

    # Select the validation mode from the active profile
    validation_mode = profile_config.get("validation_mode", "SECTIONS")
    if validation_mode == "RAW_TEXT":
        run_raw_text_validation(profile_config)
    elif validation_mode == "SECTIONS":
        run_section_validation(profile_config)
    else:
        print(f"[Error] Invalid validation_mode in profile '{active_profile_name}': '{validation_mode}'. Must be 'SECTIONS' or 'RAW_TEXT'.")


if __name__ == "__main__":
    main()
